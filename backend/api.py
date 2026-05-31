import sys, os, re
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

import json
import asyncio
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Optional
from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel

from database import get_db, init_db, SessionLocal
from database import Report, NewsItem, Angle, Headline, Schedule, RiskItem
from demo import create_mock_news, create_mock_angles, create_mock_headlines, create_mock_risks

app = FastAPI(title="NewsForge API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

ENV_PATH = Path(__file__).parent.parent / ".env"


# ── WebSocket manager ─────────────────────────────────────────────────────────

class ConnectionManager:
    def __init__(self):
        self.active: dict[str, WebSocket] = {}

    async def connect(self, ws: WebSocket, client_id: str):
        await ws.accept()
        self.active[client_id] = ws

    def disconnect(self, client_id: str):
        self.active.pop(client_id, None)

    async def send(self, client_id: str, data: dict):
        ws = self.active.get(client_id)
        if ws:
            try:
                await ws.send_json(data)
            except Exception:
                self.disconnect(client_id)


manager = ConnectionManager()


# ── Pydantic schemas ──────────────────────────────────────────────────────────

class RunRequest(BaseModel):
    geo: str
    use_mock: bool = False
    team_lead: str = ""

class FeedbackRequest(BaseModel):
    feedback: int   # -1 | 0 | 1

class TitleRequest(BaseModel):
    title: str

class ScheduleUpdate(BaseModel):
    interval_hours: float = 72.0
    enabled: bool = False

class SettingsPayload(BaseModel):
    telegram_bot_token:          Optional[str] = None
    telegram_chat_id:            Optional[str] = None
    slack_bot_token:             Optional[str] = None
    slack_channel_id:            Optional[str] = None
    news_coverage_days:          Optional[int] = None
    default_geos:                Optional[str] = None
    app_url:                     Optional[str] = None
    google_credentials_json:     Optional[str] = None  # raw JSON pasted by user → stored as base64
    google_drive_folder_id:      Optional[str] = None
    google_oauth_client_id:      Optional[str] = None
    google_oauth_client_secret:  Optional[str] = None
    google_oauth_refresh_token:  Optional[str] = None
    telegram_api_id:             Optional[str] = None
    telegram_api_hash:           Optional[str] = None
    telegram_session_string:     Optional[str] = None
    telegram_channels_ru:        Optional[str] = None
    telegram_channels_ua:        Optional[str] = None
    telegram_channels_by:        Optional[str] = None


# ── Startup ───────────────────────────────────────────────────────────────────

@app.on_event("startup")
def startup():
    init_db()
    _load_schedules()


# ── .env helpers ──────────────────────────────────────────────────────────────

def _read_env() -> dict:
    """Read env from .env file (local dev) with fallback to os.environ (production)."""
    result = {}
    # First load from os.environ (Railway / production)
    for k, v in os.environ.items():
        result[k] = v
    # Then overlay with .env file (local dev overrides)
    try:
        for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line.startswith("#") or "=" not in line:
                continue
            k, _, v = line.partition("=")
            result[k.strip()] = v.strip()
    except FileNotFoundError:
        pass
    return result


def _write_env(updates: dict) -> None:
    try:
        lines = ENV_PATH.read_text(encoding="utf-8").splitlines()
    except FileNotFoundError:
        lines = []

    updated = set()
    new_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("#") or "=" not in stripped:
            new_lines.append(line)
            continue
        key = stripped.split("=", 1)[0].strip().upper()
        if key in {k.upper() for k in updates}:
            # find original key casing
            match = next((k for k in updates if k.upper() == key), key)
            new_lines.append(f"{key}={updates[match]}")
            updated.add(key)
        else:
            new_lines.append(line)

    for k, v in updates.items():
        if k.upper() not in updated:
            new_lines.append(f"{k.upper()}={v}")

    ENV_PATH.write_text("\n".join(new_lines) + "\n", encoding="utf-8")

    # Reload into running process
    from dotenv import load_dotenv
    load_dotenv(ENV_PATH, override=True)


# ── Core pipeline ─────────────────────────────────────────────────────────────

def _get_liked_examples(geo: str, db) -> list:
    """Fetch up to 5 liked angles from previous done reports for this GEO."""
    liked = (
        db.query(Angle)
        .join(Report, Angle.report_id == Report.id)
        .filter(Report.geo == geo, Report.status == "done", Angle.feedback == 1)
        .order_by(Angle.created_at.desc())
        .limit(5)
        .all()
    )
    return [{"angle_title": a.angle_title, "offer_connection": a.offer_connection} for a in liked]


def _run_pipeline_sync(report_id: int, geo: str, use_mock: bool, team_lead: str = ""):
    """Full pipeline: parse → classify → angles (with liked examples) → headlines → risks → recs → notify."""
    db = SessionLocal()
    try:
        report = db.query(Report).filter(Report.id == report_id).first()
        report.status = "running"
        db.commit()

        if use_mock:
            news_items_raw  = create_mock_news()
            angles_raw      = create_mock_angles()
            headlines_raw   = create_mock_headlines()
            risks_raw       = create_mock_risks()
            recommendations = []
        else:
            import news_parser as np_mod
            import llm_processor as llm_mod

            parser    = np_mod.NewsParser()
            processor = llm_mod.LLMProcessor()

            # Fetch liked examples BEFORE generating new angles
            liked_examples = _get_liked_examples(geo, db)
            if liked_examples:
                print(f"[api] Using {len(liked_examples)} liked examples for {geo}")

            news_items_raw  = parser.aggregate_news(geo)

            # Telegram channels — добавляем как дополнительный источник
            try:
                from telegram_channel_parser import parse_telegram_channels, TelegramNewsItem
                from models import NewsItem as NewsItemModel
                tg_items = parse_telegram_channels(geo, days=int(os.getenv("NEWS_COVERAGE_DAYS", "7")))
                for tg in tg_items:
                    ni = NewsItemModel(
                        title            = tg.title,
                        source           = tg.source,
                        source_url       = tg.original_url or f"https://t.me/{tg.source}",
                        source_type      = "telegram",
                        category         = "",
                        description      = tg.description,
                        emotional_trigger= "",
                        urgency          = "week",
                        geo              = geo,
                        original_url     = tg.original_url,
                        date             = tg.published_at or datetime.utcnow(),
                    )
                    news_items_raw.append(ni)
                if tg_items:
                    print(f"[api] Добавлено {len(tg_items)} новостей из Telegram-каналов")
            except Exception as e:
                print(f"[api] Telegram parser error: {e}")

            news_items_raw  = processor.classify_news(news_items_raw)
            angles_raw      = processor.generate_angles(news_items_raw, geo, liked_examples=liked_examples)
            headlines_raw   = processor.generate_headlines(angles_raw)
            risks_raw       = processor.assess_risks(news_items_raw)
            recommendations = processor.generate_recommendations(angles_raw, news_items_raw)

        # Find previous report
        prev_report = (
            db.query(Report)
            .filter(Report.geo == geo, Report.status == "done", Report.id != report_id)
            .order_by(Report.id.desc())
            .first()
        )

        # Save news
        db_news: list = []
        for item in news_items_raw:
            n = NewsItem(
                report_id=report_id,
                title=item.title,
                source=item.source,
                source_type=item.source_type,
                category=item.category,
                description=item.description,
                emotional_trigger=item.emotional_trigger,
                urgency=item.urgency,
                geo=geo,
                original_url=getattr(item, "original_url", ""),
                published_at=getattr(item, "date", None),
            )
            db.add(n)
            db_news.append(n)
        db.flush()

        hash_to_db_id = {str(hash(item.title)): db_news[i].id for i, item in enumerate(news_items_raw)}

        # Save angles
        db_angles: list = []
        in_mem_to_db: dict = {}
        for a in angles_raw:
            ang = Angle(
                report_id=report_id,
                news_item_id=hash_to_db_id.get(a.news_id, 0),
                angle_title=a.angle_title,
                offer_connection=a.offer_connection,
                target_pain=a.target_pain,
                creative_type=a.creative_type,
                priority=a.priority,
            )
            db.add(ang)
            db_angles.append(ang)
        db.flush()
        for i, a in enumerate(angles_raw):
            in_mem_to_db[a.id] = db_angles[i].id

        # Save headlines
        for h in headlines_raw:
            db_angle_id = in_mem_to_db.get(h.angle_id)
            if db_angle_id is None:
                idx = h.angle_id - 1
                db_angle_id = db_angles[idx].id if 0 <= idx < len(db_angles) else db_angles[0].id
            db.add(Headline(
                report_id=report_id,
                angle_id=db_angle_id,
                text=h.text,
                format=h.format,
                character_count=h.character_count,
            ))

        # Remap recommendation angle_ids to DB ids
        for rec in recommendations:
            mem_id = rec.get("angle_id", 0)
            rec["angle_id"] = in_mem_to_db.get(mem_id, mem_id)

        urgent_count  = sum(1 for n in news_items_raw if n.urgency == "urgent_48h")
        eternal_count = sum(1 for n in news_items_raw if n.urgency == "eternal")

        report.team_lead        = team_lead
        report.prev_report_id   = prev_report.id if prev_report else None
        report.total_news       = len(news_items_raw)
        report.total_angles     = len(db_angles)
        report.total_headlines  = len(headlines_raw)
        report.recommendations  = json.dumps(recommendations, ensure_ascii=False, default=str)
        report.status           = "done"
        db.commit()

        print(f"[api] Done report_id={report_id} geo={geo} news={len(news_items_raw)} "
              f"angles={len(db_angles)} headlines={len(headlines_raw)} recs={len(recommendations)}")

        # Google Docs export
        gdocs_url = ""
        try:
            from google_docs_exporter import create_report_doc

            # Build headline map: db_angle_id → list of headline dicts
            headline_map: dict = {}
            for h in headlines_raw:
                da_id = in_mem_to_db.get(h.angle_id)
                if da_id is None:
                    idx = h.angle_id - 1
                    da_id = db_angles[idx].id if 0 <= idx < len(db_angles) else None
                if da_id:
                    headline_map.setdefault(da_id, []).append({"text": h.text, "format": h.format})

            # Build news_id → db_news mapping for risks
            hash_to_news = {str(hash(item.title)): item for item in news_items_raw}

            # Liked angles from previous report
            prev_liked = []
            if prev_report:
                prev_angles = db.query(Angle).filter(
                    Angle.report_id == prev_report.id, Angle.feedback == 1
                ).all()
                prev_liked = [{"angle_title": a.angle_title, "priority": a.priority}
                              for a in prev_angles]

            # Risks: match by news_id hash → news title, save to DB
            risks_export = []
            # Delete old risks for this report (idempotent re-run)
            db.query(RiskItem).filter(RiskItem.report_id == report_id).delete()
            for risk in risks_raw:
                matched = hash_to_news.get(risk.news_id)
                news_title = matched.title if matched else ""
                risks_export.append({
                    "news_title":               news_title,
                    "legal_risks":              risk.legal_risks,
                    "platform_ban_risk":        risk.platform_ban_risk,
                    "audience_negativity_risk": risk.audience_negativity_risk,
                    "reputation_risk":          risk.reputation_risk,
                    "expiry_date":              risk.expiry_date,
                })
                db.add(RiskItem(
                    report_id                = report_id,
                    news_title               = news_title,
                    legal_risks              = json.dumps(risk.legal_risks, ensure_ascii=False),
                    platform_ban_risk        = risk.platform_ban_risk,
                    audience_negativity_risk = risk.audience_negativity_risk,
                    reputation_risk          = risk.reputation_risk,
                    expiry_date              = risk.expiry_date or "",
                ))
            db.commit()

            # Angles with news title
            news_id_to_item = {str(hash(item.title)): item for item in news_items_raw}
            angle_to_news: dict = {}
            for a_raw in angles_raw:
                angle_to_news[a_raw.id] = news_id_to_item.get(a_raw.news_id)

            export_data = {
                "report_id":     report_id,
                "geo":           geo,
                "created_at":    report.created_at.isoformat(),
                "coverage_days": 7,
                "team_lead":     team_lead or "",
                "prev_report_id": prev_report.id if prev_report else None,
                "stats": {
                    "total_news":      len(news_items_raw),
                    "total_angles":    len(db_angles),
                    "total_headlines": len(headlines_raw),
                },
                # Блок 1 — все поля ТЗ
                "news": [
                    {
                        "title":             n.title,
                        "source":            n.source,
                        "source_type":       n.source_type,
                        "category":          n.category,
                        "description":       n.description,
                        "emotional_trigger": n.emotional_trigger,
                        "urgency":           n.urgency,
                        "original_url":      n.original_url or "",
                        "published_at":      n.published_at.isoformat() if n.published_at else "",
                    }
                    for n in db_news
                ],
                # Блок 2 — углы + привязка к инфоповоду
                "angles": [
                    {
                        "angle_title":     ang.angle_title,
                        "priority":        ang.priority,
                        "target_pain":     ang.target_pain,
                        "offer_connection": ang.offer_connection,
                        "creative_type":   ang.creative_type,
                        "news_title":      (angle_to_news.get(a_raw.id).title
                                            if angle_to_news.get(a_raw.id) else ""),
                        # Блок 3 — заголовки сгруппированы по углу
                        "headlines": [
                            {"text": h["text"], "format": h["format"],
                             "character_count": len(h["text"])}
                            for h in headline_map.get(ang.id, [])
                        ],
                    }
                    for ang, a_raw in zip(db_angles, angles_raw)
                ],
                # Блок 4
                "recommendations": recommendations,
                # Блок 5
                "risks": risks_export,
                # Блок 6 — срочность
                "urgency": {
                    "urgent":  [n.title for n in news_items_raw if n.urgency == "urgent_48h"],
                    "week":    [n.title for n in news_items_raw if n.urgency == "week"],
                    "eternal": [n.title for n in news_items_raw if n.urgency == "eternal"],
                },
                # Обратная связь по прошлому выпуску
                "prev_liked": prev_liked,
            }
            gdocs_url = create_report_doc(export_data) or ""
            if gdocs_url:
                report.gdocs_url = gdocs_url
                db.commit()
                print(f"[api] Google Docs: {gdocs_url}")
        except Exception as e:
            print(f"[api] Google Docs export error: {e}")

        # Notifications
        try:
            from notifier import notify_report_ready
            notify_report_ready(
                geo=geo, report_id=report_id,
                stats={"total_news": len(news_items_raw), "total_angles": len(db_angles),
                       "total_headlines": len(headlines_raw), "urgent_count": urgent_count,
                       "eternal_count": eternal_count},
                top_recommendations=recommendations, team_lead=team_lead,
                gdocs_url=gdocs_url,
            )
        except Exception as e:
            print(f"[api] Notification error: {e}")

    except Exception as e:
        print(f"[api] Pipeline error: {e}")
        import traceback; traceback.print_exc()
        try:
            r = db.query(Report).filter(Report.id == report_id).first()
            if r:
                r.status = "error"
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


# ── Report title generator ───────────────────────────────────────────────────

_MONTHS_RU = ["янв", "фев", "мар", "апр", "май", "июн",
               "июл", "авг", "сен", "окт", "ноя", "дек"]

def _default_title(geo: str) -> str:
    now = datetime.utcnow()
    month = _MONTHS_RU[now.month - 1]
    return f"Отчет {geo} · {now.day} {month} {now.year} {now.hour:02d}:{now.minute:02d}"


# ── Scheduler ─────────────────────────────────────────────────────────────────

try:
    from apscheduler.schedulers.background import BackgroundScheduler
    _scheduler = BackgroundScheduler()
    _scheduler.start()
    _SCHEDULER_AVAILABLE = True
except ImportError:
    _scheduler = None
    _SCHEDULER_AVAILABLE = False


def _scheduled_pipeline(geo: str):
    db = SessionLocal()
    try:
        if db.query(Report).filter(Report.geo == geo, Report.status == "running").first():
            return
        report = Report(geo=geo, status="pending", team_lead="Auto-Scheduler", title=_default_title(geo))
        db.add(report); db.commit(); db.refresh(report)
        sched = db.query(Schedule).filter(Schedule.geo == geo).first()
        if sched:
            sched.last_run = datetime.utcnow()
            sched.next_run = datetime.utcnow() + timedelta(hours=sched.interval_hours)
            db.commit()
        rid = report.id
    finally:
        db.close()
    _run_pipeline_sync(rid, geo, False, "Auto-Scheduler")


def _load_schedules():
    if not _SCHEDULER_AVAILABLE: return
    db = SessionLocal()
    try:
        for s in db.query(Schedule).filter(Schedule.enabled == True).all():
            _add_scheduler_job(s.geo, s.interval_hours)
    finally:
        db.close()


def _add_scheduler_job(geo: str, interval_hours: float):
    if not _SCHEDULER_AVAILABLE: return
    from apscheduler.triggers.interval import IntervalTrigger
    seconds = max(60, int(interval_hours * 3600))   # минимум 60 сек, всегда целое
    _scheduler.add_job(_scheduled_pipeline, trigger=IntervalTrigger(seconds=seconds),
                       args=[geo], id=f"pipeline_{geo}", replace_existing=True)


def _remove_scheduler_job(geo: str):
    if not _SCHEDULER_AVAILABLE: return
    try: _scheduler.remove_job(f"pipeline_{geo}")
    except Exception: pass


# ── Routes: reports ───────────────────────────────────────────────────────────

@app.get("/api/reports")
def list_reports(geo: Optional[str] = None, favorite: Optional[bool] = None, db: Session = Depends(get_db)):
    q = db.query(Report)
    if geo:
        q = q.filter(Report.geo == geo)
    if favorite:
        q = q.filter(Report.is_favorite == True)
    return [
        {"id": r.id, "geo": r.geo, "created_at": r.created_at.isoformat(),
         "status": r.status, "team_lead": r.team_lead or "",
         "title": r.title or "",
         "is_favorite": bool(r.is_favorite),
         "total_news": r.total_news, "total_angles": r.total_angles, "total_headlines": r.total_headlines}
        for r in q.order_by(Report.created_at.desc()).limit(50).all()
    ]


@app.patch("/api/reports/{report_id}/favorite")
def toggle_favorite(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")
    report.is_favorite = not bool(report.is_favorite)
    db.commit()
    return {"id": report.id, "is_favorite": bool(report.is_favorite)}


@app.patch("/api/reports/{report_id}/title")
def update_title(report_id: int, body: TitleRequest, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")
    report.title = body.title.strip()[:500]
    db.commit()
    return {"id": report.id, "title": report.title}


@app.get("/api/reports/{report_id}")
def get_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    news      = db.query(NewsItem).filter(NewsItem.report_id == report_id).all()
    angles    = db.query(Angle).filter(Angle.report_id == report_id).all()
    headlines = db.query(Headline).filter(Headline.report_id == report_id).all()
    news_by_id = {n.id: n for n in news}

    try:
        recs = json.loads(report.recommendations) if report.recommendations else []
    except Exception:
        recs = []

    prev_liked = []
    if report.prev_report_id:
        prev_angles = db.query(Angle).filter(
            Angle.report_id == report.prev_report_id, Angle.feedback == 1).all()
        prev_liked = [{"id": a.id, "angle_title": a.angle_title, "priority": a.priority}
                      for a in prev_angles]

    risks_db = db.query(RiskItem).filter(RiskItem.report_id == report_id).all()

    return {
        "id": report.id, "geo": report.geo,
        "created_at": report.created_at.isoformat(),
        "status": report.status, "team_lead": report.team_lead or "",
        "title": report.title or "",
        "prev_report_id": report.prev_report_id, "coverage_days": 7,
        "gdocs_url": report.gdocs_url or "",
        "is_favorite": bool(report.is_favorite),
        "stats": {"total_news": report.total_news, "total_angles": report.total_angles,
                  "total_headlines": report.total_headlines},
        "recommendations": recs,
        "prev_performance": {"liked": prev_liked, "count": len(prev_liked)},
        "news": [
            {"id": n.id, "title": n.title, "source": n.source, "source_type": n.source_type,
             "category": n.category, "description": n.description,
             "emotional_trigger": n.emotional_trigger, "urgency": n.urgency,
             "original_url": n.original_url,
             "published_at": n.published_at.isoformat() if n.published_at else None}
            for n in news
        ],
        "angles": [
            {"id": a.id, "news_item_id": a.news_item_id,
             "news_title": news_by_id[a.news_item_id].title if a.news_item_id in news_by_id else "",
             "angle_title": a.angle_title, "offer_connection": a.offer_connection,
             "target_pain": a.target_pain, "creative_type": a.creative_type,
             "priority": a.priority, "feedback": a.feedback,
             "headlines": [{"id": h.id, "text": h.text, "format": h.format,
                            "character_count": h.character_count}
                           for h in headlines if h.angle_id == a.id]}
            for a in angles
        ],
        "risks": [
            {"id": r.id,
             "news_title":               r.news_title,
             "legal_risks":              json.loads(r.legal_risks) if r.legal_risks else [],
             "platform_ban_risk":        r.platform_ban_risk,
             "audience_negativity_risk": r.audience_negativity_risk,
             "reputation_risk":          r.reputation_risk,
             "expiry_date":              r.expiry_date}
            for r in risks_db
        ],
    }


@app.post("/api/reports/{report_id}/export-gdocs")
def export_report_gdocs(report_id: int, db: Session = Depends(get_db)):
    """Re-export an existing report to Google Docs (useful for testing or re-generation)."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")

    news    = db.query(NewsItem).filter(NewsItem.report_id == report_id).all()
    angles  = db.query(Angle).filter(Angle.report_id == report_id).all()
    headlines = db.query(Headline).filter(Headline.report_id == report_id).all()

    news_by_id = {n.id: n for n in news}
    headline_map: dict = {}
    for h in headlines:
        headline_map.setdefault(h.angle_id, []).append({"text": h.text, "format": h.format})

    prev_liked = []
    if report.prev_report_id:
        prev_angles = db.query(Angle).filter(
            Angle.report_id == report.prev_report_id, Angle.feedback == 1
        ).all()
        prev_liked = [{"angle_title": a.angle_title, "priority": a.priority} for a in prev_angles]

    try:
        recs = json.loads(report.recommendations) if report.recommendations else []
    except Exception:
        recs = []

    export_data = {
        "report_id":    report_id,
        "geo":          report.geo,
        "created_at":   report.created_at.isoformat(),
        "coverage_days": 7,
        "team_lead":    report.team_lead or "",
        "prev_report_id": report.prev_report_id,
        "stats": {
            "total_news":      report.total_news,
            "total_angles":    report.total_angles,
            "total_headlines": report.total_headlines,
        },
        "news": [
            {
                "title":             n.title,
                "source":            n.source,
                "source_type":       n.source_type,
                "category":          n.category,
                "description":       n.description,
                "emotional_trigger": n.emotional_trigger,
                "urgency":           n.urgency,
                "original_url":      n.original_url or "",
                "published_at":      n.published_at.isoformat() if n.published_at else "",
            }
            for n in news
        ],
        "angles": [
            {
                "angle_title":     a.angle_title,
                "priority":        a.priority,
                "target_pain":     a.target_pain,
                "offer_connection": a.offer_connection,
                "creative_type":   a.creative_type,
                "news_title":      news_by_id[a.news_item_id].title if a.news_item_id in news_by_id else "",
                "headlines": [
                    {"text": h["text"], "format": h["format"], "character_count": len(h["text"])}
                    for h in headline_map.get(a.id, [])
                ],
            }
            for a in angles
        ],
        "recommendations": recs,
        "risks": [
            {
                "news_title":               r.news_title,
                "legal_risks":              json.loads(r.legal_risks) if r.legal_risks else [],
                "platform_ban_risk":        r.platform_ban_risk,
                "audience_negativity_risk": r.audience_negativity_risk,
                "reputation_risk":          r.reputation_risk,
                "expiry_date":              r.expiry_date,
            }
            for r in db.query(RiskItem).filter(RiskItem.report_id == report_id).all()
        ],
        "urgency": {
            "urgent":  [n.title for n in news if n.urgency == "urgent_48h"],
            "week":    [n.title for n in news if n.urgency == "week"],
            "eternal": [n.title for n in news if n.urgency == "eternal"],
        },
        "prev_liked": prev_liked,
    }

    try:
        from google_docs_exporter import create_report_doc
        gdocs_url = create_report_doc(export_data) or ""
        if gdocs_url:
            report.gdocs_url = gdocs_url
            db.commit()
        return {"ok": True, "gdocs_url": gdocs_url}
    except Exception as e:
        raise HTTPException(500, f"Google Docs export failed: {e}")


@app.delete("/api/reports/{report_id}")
def delete_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "Report not found")
    db.query(Headline).filter(Headline.report_id == report_id).delete()
    db.query(Angle).filter(Angle.report_id == report_id).delete()
    db.query(NewsItem).filter(NewsItem.report_id == report_id).delete()
    db.query(RiskItem).filter(RiskItem.report_id == report_id).delete()
    db.query(Report).filter(Report.id == report_id).delete()
    db.commit()
    return {"ok": True}


# ── Routes: run ───────────────────────────────────────────────────────────────

@app.post("/api/run")
async def trigger_run(req: RunRequest, background_tasks: BackgroundTasks,
                      db: Session = Depends(get_db)):
    report = Report(geo=req.geo, status="pending", team_lead=req.team_lead, title=_default_title(req.geo))
    db.add(report); db.commit(); db.refresh(report)
    client_id = f"{req.geo}-{report.id}"
    background_tasks.add_task(_run_pipeline_sync, report.id, req.geo, req.use_mock, req.team_lead)
    return {"report_id": report.id, "client_id": client_id}


@app.websocket("/ws/{client_id}")
async def ws_endpoint(websocket: WebSocket, client_id: str):
    await manager.connect(websocket, client_id)
    try:
        data = await websocket.receive_json()
        db = SessionLocal()
        try:
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None, _run_pipeline_sync,
                data.get("report_id"), data.get("geo", "RU"),
                data.get("use_mock", False), data.get("team_lead", "")
            )
            await manager.send(client_id, {"done": True, "report_id": data.get("report_id")})
        finally:
            db.close()
    except WebSocketDisconnect:
        manager.disconnect(client_id)


# ── Routes: feedback ──────────────────────────────────────────────────────────

@app.patch("/api/angles/{angle_id}/feedback")
def set_feedback(angle_id: int, req: FeedbackRequest, db: Session = Depends(get_db)):
    angle = db.query(Angle).filter(Angle.id == angle_id).first()
    if not angle:
        raise HTTPException(404)
    angle.feedback = req.feedback
    db.commit()
    return {"ok": True}


# ── Routes: schedule ──────────────────────────────────────────────────────────

@app.get("/api/schedule")
def list_schedule(db: Session = Depends(get_db)):
    rows = {s.geo: s for s in db.query(Schedule).all()}
    return [
        {"geo": geo, "interval_hours": rows[geo].interval_hours if geo in rows else 72,
         "enabled": rows[geo].enabled if geo in rows else False,
         "last_run": rows[geo].last_run.isoformat() if geo in rows and rows[geo].last_run else None,
         "next_run": rows[geo].next_run.isoformat() if geo in rows and rows[geo].next_run else None,
         "scheduler_available": _SCHEDULER_AVAILABLE}
        for geo in ["RU", "UA", "BY", "KZ", "DE", "PL"]
    ]


@app.put("/api/schedule/{geo}")
def update_schedule(geo: str, req: ScheduleUpdate, db: Session = Depends(get_db)):
    s = db.query(Schedule).filter(Schedule.geo == geo).first()
    if not s:
        s = Schedule(geo=geo); db.add(s)
    s.interval_hours = req.interval_hours
    s.enabled = req.enabled
    if req.enabled:
        s.next_run = datetime.utcnow() + timedelta(hours=req.interval_hours)
        _add_scheduler_job(geo, req.interval_hours)
    else:
        _remove_scheduler_job(geo); s.next_run = None
    db.commit()
    return {"geo": s.geo, "interval_hours": s.interval_hours, "enabled": s.enabled,
            "next_run": s.next_run.isoformat() if s.next_run else None}


# ── Routes: export ───────────────────────────────────────────────────────────

@app.get("/api/reports/{report_id}/export.docx")
def export_docx(report_id: int, db: Session = Depends(get_db)):
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    report = db.query(Report).filter(Report.id == report_id).first()
    if not report: raise HTTPException(404, "Report not found")

    news      = db.query(NewsItem).filter(NewsItem.report_id == report_id).all()
    angles    = db.query(Angle).filter(Angle.report_id == report_id).all()
    headlines = db.query(Headline).filter(Headline.report_id == report_id).all()
    try:
        recs = json.loads(report.recommendations) if report.recommendations else []
    except Exception:
        recs = []

    doc = Document()

    # Styles helper
    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        return p

    def para(text, bold=False, italic=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading(f"NewsForge — Отчёт {report.geo} #{report_id}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    created = report.created_at.strftime("%d.%m.%Y %H:%M") if report.created_at else ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Дата: {created}").italic = True
    if report.team_lead:
        p.add_run(f"  |  Тимлид: {report.team_lead}").italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        f"Инфоповодов: {report.total_news}  |  "
        f"Углов: {report.total_angles}  |  "
        f"Заголовков: {report.total_headlines}"
    ).bold = True

    doc.add_paragraph()

    # ── Рекомендации ───────────────────────────────────────────────────────────
    if recs:
        heading("Топ-5 рекомендаций к тесту")
        for rec in recs[:5]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(rec.get("angle_title", "")).bold = True
            if rec.get("reasoning"):
                doc.add_paragraph(f"   {rec['reasoning']}")
        doc.add_paragraph()

    # ── Инфоповоды ─────────────────────────────────────────────────────────────
    heading(f"Инфоповоды ({len(news)})")
    for i, n in enumerate(news, 1):
        urgency_mark = "🔥 " if n.urgency == "urgent_48h" else "♾ " if n.urgency == "eternal" else ""
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{urgency_mark}{n.title}").bold = True
        pub = n.published_at.strftime("%d.%m.%Y") if n.published_at else ""
        src_line = f"   Источник: {n.source} · {n.source_type or ''}"
        if pub:
            src_line += f" · {pub}"
        doc.add_paragraph(src_line)
    doc.add_paragraph()

    # ── Углы и заголовки ───────────────────────────────────────────────────────
    heading(f"Маркетинговые углы ({len(angles)})")
    hl_by_angle = {}
    for h in headlines:
        hl_by_angle.setdefault(h.angle_id, []).append(h)

    for ang in angles:
        p = doc.add_heading(f"[{ang.priority}] {ang.angle_title}", level=2)
        if ang.target_pain:
            doc.add_paragraph(f"Боль аудитории: {ang.target_pain}")
        if ang.offer_connection:
            doc.add_paragraph(f"Связь с оффером: {ang.offer_connection}")
        if ang.creative_type:
            doc.add_paragraph(f"Тип: {ang.creative_type}")
        hl_list = hl_by_angle.get(ang.id, [])
        if hl_list:
            doc.add_paragraph("Заголовки:").runs[0].bold = True
            for h in hl_list:
                p = doc.add_paragraph(h.text, style="List Bullet")
        doc.add_paragraph()

    # ── Сохранить ──────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    geo_safe = report.geo.upper()
    filename = f"newsforge-{geo_safe}-{report_id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Routes: settings ─────────────────────────────────────────────────────────

@app.get("/api/settings")
def get_settings():
    env = _read_env()
    def configured(key): return bool(env.get(key, "").strip())
    gdocs_ok = configured("GOOGLE_CREDENTIALS_B64") or (
        bool(env.get("GOOGLE_SERVICE_ACCOUNT_PATH", "").strip()) and
        os.path.exists(env.get("GOOGLE_SERVICE_ACCOUNT_PATH", ""))
    )
    folder_id = env.get("GOOGLE_DRIVE_FOLDER_ID", "").strip()
    # Get SA email for UI display
    sa_email = ""
    if gdocs_ok:
        try:
            import base64 as _b64, json as _json
            b64 = env.get("GOOGLE_CREDENTIALS_B64", "")
            if b64:
                info = _json.loads(_b64.b64decode(b64).decode("utf-8"))
                sa_email = info.get("client_email", "")
        except Exception:
            pass
    oauth_ok = (
        configured("GOOGLE_OAUTH_REFRESH_TOKEN") and
        configured("GOOGLE_OAUTH_CLIENT_ID") and
        configured("GOOGLE_OAUTH_CLIENT_SECRET")
    )
    return {
        "anthropic_configured":    configured("ANTHROPIC_API_KEY"),
        "telegram_configured":     configured("TELEGRAM_BOT_TOKEN"),
        "telegram_chat_id":        env.get("TELEGRAM_CHAT_ID", ""),
        "slack_configured":        configured("SLACK_BOT_TOKEN"),
        "slack_channel_id":        env.get("SLACK_CHANNEL_ID", ""),
        "news_coverage_days":      int(env.get("NEWS_COVERAGE_DAYS", "7")),
        "default_geos":            env.get("DEFAULT_GEOS", "RU,UA,BY"),
        "app_url":                 env.get("APP_URL", "http://localhost:5173"),
        "gdocs_configured":        gdocs_ok,
        "gdocs_folder_configured": bool(folder_id),
        "google_drive_folder_id":  folder_id,
        "sa_email":                sa_email,
        "gdocs_oauth_configured":  oauth_ok,
        "gdocs_auth_mode":         "oauth" if oauth_ok else ("service_account" if gdocs_ok else "none"),
        "telegram_channels_configured": (
            configured("TELEGRAM_API_ID") and
            configured("TELEGRAM_API_HASH") and
            configured("TELEGRAM_SESSION_STRING")
        ),
    }


@app.post("/api/settings")
def save_settings(payload: SettingsPayload):
    import base64 as _b64
    updates = {}
    if payload.telegram_bot_token      is not None:
        updates["TELEGRAM_BOT_TOKEN"] = payload.telegram_bot_token
    if payload.telegram_chat_id        is not None:
        updates["TELEGRAM_CHAT_ID"] = payload.telegram_chat_id
    if payload.slack_bot_token         is not None:
        updates["SLACK_BOT_TOKEN"] = payload.slack_bot_token
    if payload.slack_channel_id        is not None:
        updates["SLACK_CHANNEL_ID"] = payload.slack_channel_id
    if payload.news_coverage_days      is not None:
        updates["NEWS_COVERAGE_DAYS"] = str(payload.news_coverage_days)
    if payload.default_geos            is not None:
        updates["DEFAULT_GEOS"] = payload.default_geos
    if payload.app_url                 is not None:
        updates["APP_URL"] = payload.app_url.rstrip("/")
    if payload.google_credentials_json is not None and payload.google_credentials_json.strip():
        try:
            json.loads(payload.google_credentials_json)
            encoded = _b64.b64encode(payload.google_credentials_json.encode("utf-8")).decode("ascii")
            updates["GOOGLE_CREDENTIALS_B64"] = encoded
        except Exception:
            raise HTTPException(400, "google_credentials_json is not valid JSON")
    if payload.google_drive_folder_id is not None:
        updates["GOOGLE_DRIVE_FOLDER_ID"] = payload.google_drive_folder_id
    if payload.google_oauth_client_id is not None:
        updates["GOOGLE_OAUTH_CLIENT_ID"] = payload.google_oauth_client_id
    if payload.google_oauth_client_secret is not None:
        updates["GOOGLE_OAUTH_CLIENT_SECRET"] = payload.google_oauth_client_secret
    if payload.google_oauth_refresh_token is not None:
        updates["GOOGLE_OAUTH_REFRESH_TOKEN"] = payload.google_oauth_refresh_token
    if payload.telegram_api_id is not None:
        updates["TELEGRAM_API_ID"] = payload.telegram_api_id
    if payload.telegram_api_hash is not None:
        updates["TELEGRAM_API_HASH"] = payload.telegram_api_hash
    if payload.telegram_session_string is not None:
        updates["TELEGRAM_SESSION_STRING"] = payload.telegram_session_string
    if payload.telegram_channels_ru is not None:
        updates["TELEGRAM_CHANNELS_RU"] = payload.telegram_channels_ru
    if payload.telegram_channels_ua is not None:
        updates["TELEGRAM_CHANNELS_UA"] = payload.telegram_channels_ua
    if payload.telegram_channels_by is not None:
        updates["TELEGRAM_CHANNELS_BY"] = payload.telegram_channels_by

    if updates:
        _write_env(updates)
    return {"ok": True, "updated": list(updates.keys())}


@app.post("/api/settings/test-notify")
def test_notify():
    from notifier import send_telegram, send_slack
    tg = send_telegram("🧪 <b>NewsForge</b> — тестовое уведомление работает!")
    sl = send_slack("🧪 NewsForge — test notification works!")
    return {"telegram": tg, "slack": sl}


# ── Routes: stats ─────────────────────────────────────────────────────────────

@app.get("/api/stats")
def stats(db: Session = Depends(get_db)):
    return {
        "total_reports":   db.query(Report).filter(Report.status == "done").count(),
        "total_angles":    db.query(Angle).count(),
        "total_headlines": db.query(Headline).count(),
        "liked_angles":    db.query(Angle).filter(Angle.feedback == 1).count(),
        "disliked_angles": db.query(Angle).filter(Angle.feedback == -1).count(),
    }


# ── Static frontend (production) ──────────────────────────────────────────────
# Serve the built React app for all non-API routes

_STATIC_DIR = Path(__file__).parent.parent / "frontend" / "dist"

if _STATIC_DIR.exists():
    app.mount("/assets", StaticFiles(directory=_STATIC_DIR / "assets"), name="assets")

    @app.get("/{full_path:path}", include_in_schema=False)
    async def serve_spa(full_path: str):
        """Serve React SPA — return index.html for all non-API routes."""
        index = _STATIC_DIR / "index.html"
        return FileResponse(index)
